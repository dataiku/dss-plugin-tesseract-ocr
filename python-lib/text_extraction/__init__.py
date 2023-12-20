import docx
from io import BytesIO
import logging
import os
import pypandoc
import pypdfium2 as pdfium
import tempfile


logger = logging.getLogger(__name__)


def download_pandoc_binaries():
    """ download pandoc prebuilt binaries into the temporary job folder """
    try:
        import pypandoc
        pandoc_tmp_directory = os.getcwd()
        pypandoc.ensure_pandoc_installed(targetfolder=pandoc_tmp_directory)
        return True
    except Exception as e:
        logger.warning("Failed to download pandoc binaries: {}".format(e))
        return False


def extract_text_content(file_bytes, extension, with_pandoc):
    """
    Extract text content from file bytes:
    - First try to extract from PDF or docx file
    - Then try using pandoc to extract other files into plain text.
    - Finally, just decode the bytes if pandoc failed or is not downloaded.
    """
    if extension == "doc":
        raise ValueError("'doc' files are not supported, try to convert them to docx.")

    if extension == "pdf":
        pdf_pages = pdfium.PdfDocument(file_bytes)
        return "\n".join([page.get_textpage().get_text_range() for page in pdf_pages])
    elif extension == "docx":
        doc = docx.Document(BytesIO(file_bytes))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        text = ""
        if with_pandoc:
            try:
                temporary_job_folder = os.getcwd()
                with tempfile.NamedTemporaryFile(dir=temporary_job_folder, suffix=".{}".format(extension)) as tmp:
                    tmp.write(file_bytes)
                    text = pypandoc.convert_file(tmp.name, to="plain", format=extension)

            except Exception as e:
                logger.warning("Failed to extract text with pandoc: {}".format(e))

        # pandoc outputs empty string with new line for some formats that work with .decode() (for instance csv/tsv)
        if not text.strip():
            return file_bytes.decode()
        
        return text


def extract_text_chunks(filename, file_bytes, extension, with_pandoc, metadata_as_plain_text, use_pdf_bookmarks):
    if extension == "doc":
        raise ValueError("'doc' files are not supported, try to convert them to docx.")

    try:
        if extension == "pdf":
            pdf_pages = pdfium.PdfDocument(file_bytes)
            outline = list(pdf_pages.get_toc())
            if len(outline) == 0 or not use_pdf_bookmarks:
                # only extract page numbers when no outline is found
                return [
                    {
                        'file': filename,
                        'text': page.get_textpage().get_text_range(),
                        'chunk_id': page_id + 1,
                        'metadata': "Page {}".format(page_id + 1) if metadata_as_plain_text else {"page": page_id + 1},
                        'error_message': ""
                    }
                    for page_id, page in enumerate(pdf_pages)
                ]
            else:
                return _extract_pdf_chunks(filename, pdf_pages, outline, metadata_as_plain_text)
        elif extension == "md":
            return _extract_markdown_chunks(file_bytes.decode(), filename, metadata_as_plain_text)
        else:
            if not with_pandoc:
                raise ValueError("pandoc is required to extract chunks from files (except for PDF and markdown).")
            temporary_job_folder = os.getcwd()
            with tempfile.NamedTemporaryFile(dir=temporary_job_folder, suffix=".{}".format(extension)) as tmp:
                tmp.write(file_bytes)
                # 'gfm' is for markdown_github, a simplified form of markdown for more consistent results across OSes
                markdown = pypandoc.convert_file(tmp.name, to="gfm", format=extension)

                if not markdown.strip():
                    raise ValueError("Content is empty after converting to markdown.")

                return _extract_markdown_chunks(markdown, filename, metadata_as_plain_text, convert_text_blocks=True, markdown_format="gfm")
    except Exception as e:
        logger.warning("Failed to extract chunks, falling back to text content extraction: {}".format(e))
        text = extract_text_content(file_bytes, extension, with_pandoc)
        return [{
            'file': filename,
            'text': text,
            'chunk_id': 1,
            'metadata': "",
            'error_message': "Failed to extract chunks, fallback to text content extraction"
        }]


def _extract_markdown_chunks(markdown, filename, metadata_as_plain_text, convert_text_blocks=False, markdown_format="gfm"):
    """
    Extracts chunks from a markdown document. 
    
    The document is chunked according to its headers. Headers with levels 4+ are considered
    as part of the text.

    Returns a list of dictionaries, where each dictionary contains a chunk data and 
    corresponding metadata.

    This code is largely inspired by `langchain.text_splitter.MarkdownHeaderTextSplitter.split_text`
    https://github.com/langchain-ai/langchain/blob/v0.0.333/libs/langchain/langchain/text_splitter.py#L376
    """

    lines = markdown.split("\n")
    # Final output
    lines_with_metadata = []
    # text and metadata of the chunk currently being processed
    current_text = []
    current_metadata = {}
    # Keep track of the nested header structure
    # header_stack: List[Dict[str, Union[int, str]]] = []
    header_stack = []
    initial_metadata = {}

    in_code_block = False

    for line in lines:
        stripped_line = line
        if not stripped_line.startswith(("    ", "\t")):
            # Header lines can only start with 0 to 3 spaces
            stripped_line = line.lstrip()

        if stripped_line.startswith(("```", "~~~")):
            # code block in one row (only for ```) 
            if stripped_line.count("```") >= 2:
                in_code_block = False
            else:
                in_code_block = not in_code_block

        if in_code_block:
            current_text.append(line)
            continue

        # Check each line against each of the header types (e.g., #, ##), header_level is the number of '#'
        for header_level in range(1, 4):
            # Check if line starts with a header that we intend to split on
            if stripped_line.startswith("#" * header_level) and (
                # Header with no text OR header is followed by space are valid conditions that sep is being used a header
                len(stripped_line) == header_level or stripped_line[header_level] == " "
            ):

                # Pop out headers of lower or same level from the stack
                while header_stack and header_stack[-1]["level"] >= header_level:
                    # We have encountered a new header at the same or higher level
                    popped_header = header_stack.pop()
                    # Clear the metadata for the popped header in initial_metadata
                    if popped_header["level"] in initial_metadata:
                        initial_metadata.pop(popped_header["level"])

                # Push the current header to the stack
                if convert_text_blocks:
                    data = pypandoc.convert_text(stripped_line, to="plain", format=markdown_format).strip()
                else:
                    data = stripped_line[header_level :].strip()
                header = {"level": header_level, "data": data}
                header_stack.append(header)
                # Update initial_metadata with the current header
                initial_metadata[header_level] = header["data"]

                # Add the previous line to the lines_with_metadata only if current_text is not empty
                if current_text:
                    if any(current_text):  # Add only chunks that contain at least one non-empty element
                        if convert_text_blocks:
                            text = pypandoc.convert_text("\n".join(current_text), to="plain", format=markdown_format).strip()
                        else:
                            text = "\n".join(current_text)
                        lines_with_metadata.append({"text": text, "metadata": current_metadata.copy()})
                    current_text.clear()

                break
        else:
            # Add line to current_text when no header was found
            current_text.append(line)

        current_metadata = initial_metadata.copy()

    if current_text:
        if convert_text_blocks:
            text = pypandoc.convert_text("\n".join(current_text), to="plain", format=markdown_format).strip()
        else:
            text = "\n".join(current_text)
        lines_with_metadata.append({"text": text, "metadata": current_metadata})

    chunks = []
    for line_id, line in enumerate(lines_with_metadata):
        # Header metadata is encoded as {"headers": ["header 1", "header 2", "header 3"]} or "header 1 > header 2 > header 3" in plain text
        headers = list(line["metadata"].values())
        header_metadata = ""
        if len(headers) > 0:
            header_metadata = " > ".join(headers) if metadata_as_plain_text else {"headers": headers}
        chunks.append({"file": filename, "text": line["text"], "chunk_id": line_id + 1, "metadata": header_metadata, "error_message": ""})

    return chunks


def _extract_text_from_pdf_bound(pdf_pages, start_page, start_vertical_position, end_page, end_vertical_position):
    """
    Extract text between a starting vertical position in a starting page and an ending vertical position in an ending page.
    """
    text = ""
    while start_page < end_page:
        text += pdf_pages[start_page].get_textpage().get_text_bounded(top=start_vertical_position)
        start_page += 1
        start_vertical_position = None
    text += pdf_pages[start_page].get_textpage().get_text_bounded(top=start_vertical_position, bottom=end_vertical_position)
    return text


def _extract_pdf_chunks(filename, pdf_pages, outline, metadata_as_plain_text):
    """
    Extract chunks from a PDF outline
    """
    chunks = []
    headers = []

    # extract text before first header and add it only if the text is not empty
    text = _extract_text_from_pdf_bound(pdf_pages, start_page=0, start_vertical_position=None, end_page=0, end_vertical_position=outline[0].view_pos[1])
    if text.strip():
        chunks.append({"file": filename, "text": text, "chunk_id": 1, "metadata": "", "error_message": ""})

    for outline_id in range(len(outline)):
        title = outline[outline_id].title
        level = outline[outline_id].level
        start_page = outline[outline_id].page_index
        start_vertical_position = outline[outline_id].view_pos[1]

        # for the last header, extract until the end of the PDF file
        last_header = (outline_id == len(outline) - 1)
        end_page = outline[outline_id+1].page_index if not last_header else len(pdf_pages)-1
        end_vertical_position = outline[outline_id+1].view_pos[1] if not last_header else None

        # headers contains the list of the current parent headers and the current header 
        if len(headers) < level + 1:  # going into a child header
            headers.append(title)
        elif len(headers) == level + 1:  # staying at the same header level
            headers[level] = title
        else:  # going back to a higher header level
            headers = headers[:level] + [title]

        text = _extract_text_from_pdf_bound(pdf_pages, start_page, start_vertical_position, end_page, end_vertical_position)

        header_metadata = " > ".join(headers) if metadata_as_plain_text else {"headers": headers.copy()}
        chunks.append({"file": filename, "text": text, "chunk_id": len(chunks) + 1, "metadata": header_metadata, "error_message": ""})

    return chunks
