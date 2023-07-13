import docx
from io import BytesIO
import logging
from ocr_constants import Constants
import os
import pypandoc
import pypdfium2 as pdfium
from shutil import which
import tempfile


logger = logging.getLogger(__name__)


def pdf_to_pil_images_iterator(pdf_bytes, dpi=None):
    """ iterator over the multiple images of pdf bytes """
    pdf_pages = pdfium.PdfDocument(pdf_bytes)
    # scale is DPI / 72 according to pypdfium2 doc
    scale = dpi / 72 if dpi else 2
    for pdf_page in pdf_pages:
        yield pdf_page.render(scale=scale).to_pil()


def convert_image_to_greyscale_bytes(img, quality=75):
    """ convert a PIL image to greyscale with a specified dpi and output image as bytes """
    img = img.convert('L')
    buf = BytesIO()
    img.save(buf, format='JPEG', quality=quality)
    return buf.getvalue()


def image_conversion_parameters(recipe_config):
    """ retrieve image conversion recipe parameters """
    params = {}
    dpi = recipe_config.get(Constants.DPI, 200)
    assert dpi > 0 and dpi <= 4000
    params[Constants.DPI] = dpi

    quality = recipe_config.get(Constants.QUALITY, 75)
    assert quality > 0 and quality <= 95
    params[Constants.QUALITY] = quality

    return params


def image_processing_parameters(recipe_config):
    """ retrieve image processing recipe parameters """
    params = {}
    params[Constants.FUNCTIONS_DEF] = recipe_config.get(Constants.FUNCTIONS_DEF, None)
    params[Constants.PIPELINE_DEF] = recipe_config.get(Constants.PIPELINE_DEF, None)
    return params


def ocr_text_extraction_parameters(recipe_config):
    """ retrieve text extraction recipe parameters """
    params = {}
    params[Constants.RECOMBINE_PDF] = recipe_config.get(Constants.RECOMBINE_PDF, False)
    selected_ocr_engine = recipe_config.get(Constants.OCR_ENGINE, Constants.DEFAULT_ENGINE)
    advanced = recipe_config.get('advanced_parameters', False)

    if selected_ocr_engine == Constants.DEFAULT_ENGINE:
        advanced = False
        selected_ocr_engine = get_default_ocr_engine()

    params[Constants.OCR_ENGINE] = selected_ocr_engine

    if params[Constants.OCR_ENGINE] == Constants.TESSERACT:
        params[Constants.LANGUAGE_TESSERACT] = recipe_config.get(Constants.LANGUAGE_TESSERACT, "eng") if advanced else "eng"
    elif params[Constants.OCR_ENGINE] == Constants.EASYOCR:
        import easyocr
        language = recipe_config.get(Constants.LANGUAGE_EASYOCR, "en") if advanced else "en"
        # instantiate the easyocr.Reader only once here because it takes some time
        # use tmp folders inside the job temporary folder to store the model and the custom network model (note that this one isn't used)
        model_storage_directory = os.path.join(os.getcwd(), "easyocr_model_tmp")
        user_network_directory = os.path.join(os.getcwd(), "easyocr_user_network_tmp")
        params[Constants.EASYOCR_READER] = easyocr.Reader(
            lang_list=[language], gpu=False,
            model_storage_directory=model_storage_directory,
            user_network_directory=user_network_directory,
            verbose=False
        )

    return params


def get_default_ocr_engine():
    if which("tesseract") is not None:  # check if tesseract is in the path
        return Constants.TESSERACT
    return Constants.EASYOCR


def try_download_pandoc():
    """ download pandoc prebuilt binaries into the temporary job folder """
    try:
        import pypandoc
        pandoc_tmp_directory = os.getcwd()
        pypandoc.download_pandoc(targetfolder=pandoc_tmp_directory, download_folder=pandoc_tmp_directory)
        return True
    except Exception as e:
        logger.warning("Failed to download pandoc binaries: {}".format(e))
        return False


def extract_text(file_bytes, extension, with_pandoc):
    """
    First extract with pdfium or docx PDFs and docx files.
    Then try to extract text using pandoc into plain text.
    Finally, just decode the bytes if pandoc failed or is not downloaded.
    """
    if extension == "pdf":
        pdf_pages = pdfium.PdfDocument(file_bytes)
        return "\n".join([page.get_textpage().get_text_range() for page in pdf_pages])
    elif extension == "docx":
        doc = docx.Document(BytesIO(file_bytes))
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    else:
        text = ""
        if with_pandoc:
            try:
                temporary_job_folder = os.getcwd()
                with tempfile.NamedTemporaryFile(dir=temporary_job_folder, suffix=".{}".format(extension)) as tmp:
                    tmp.write(file_bytes)
                    text = pypandoc.convert_file(tmp.name, to="plain", format=extension)

            except Exception as e:
                logger.warning("Failed to extract text with pandoc: {}".format(e))

        if not text.strip():
            return file_bytes.decode()
        
        return text
            
        
